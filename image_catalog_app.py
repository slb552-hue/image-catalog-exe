import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from PIL import Image, ImageTk

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader


APP_TITLE = "קטלוג תמונות - 2 בעמוד (PDF / Word)"


def make_caption(x, y):
    return f"תמונה {x}- מבנה מס׳ {y}"


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE)
        self.geometry("1100x650")

        self.items = []  # {path, x, y}

        self._build_ui()

    def _build_ui(self):
        top = ttk.Frame(self)
        top.pack(fill="x", padx=10, pady=10)

        ttk.Button(top, text="הוסף תמונות", command=self.add_images).pack(side="left")
        ttk.Button(top, text="מחק נבחר", command=self.delete_selected).pack(side="left", padx=(10, 0))

        ttk.Separator(top, orient="vertical").pack(side="left", fill="y", padx=12)

        ttk.Button(top, text="ייצא ל-PDF", command=self.export_pdf).pack(side="left")
        ttk.Button(top, text="ייצא ל-Word (DOCX)", command=self.export_docx).pack(side="left", padx=(10, 0))

        main = ttk.Frame(self)
        main.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        left = ttk.Frame(main)
        left.pack(side="left", fill="both", expand=True)

        cols = ("file", "x", "y")
        self.tree = ttk.Treeview(left, columns=cols, show="headings", selectmode="browse")
        self.tree.heading("file", text="קובץ")
        self.tree.heading("x", text="תמונה X")
        self.tree.heading("y", text="מבנה Y")

        self.tree.column("file", width=520)
        self.tree.column("x", width=90, anchor="center")
        self.tree.column("y", width=90, anchor="center")

        self.tree.pack(fill="both", expand=True)
        self.tree.bind("<<TreeviewSelect>>", self.on_select)

        edit = ttk.Frame(left)
        edit.pack(fill="x", pady=8)

        ttk.Label(edit, text="תמונה X:").pack(side="left")
        self.x_var = tk.StringVar()
        ttk.Entry(edit, width=10, textvariable=self.x_var).pack(side="left", padx=(5, 15))

        ttk.Label(edit, text="מבנה Y:").pack(side="left")
        self.y_var = tk.StringVar()
        ttk.Entry(edit, width=10, textvariable=self.y_var).pack(side="left", padx=(5, 15))

        ttk.Button(edit, text="שמור ערכים לשורה", command=self.save_xy).pack(side="left")

        right = ttk.Frame(main, width=350)
        right.pack(side="right", fill="y", padx=(10, 0))

        ttk.Label(right, text="תצוגה מקדימה").pack(anchor="center", pady=(0, 6))
        self.preview_label = ttk.Label(right)
        self.preview_label.pack(fill="both", expand=True)

        self.caption_label = ttk.Label(right, text="", anchor="e", justify="right", font=("Arial", 12))
        self.caption_label.pack(fill="x", pady=8)

    def add_images(self):
        paths = filedialog.askopenfilenames(
            title="בחר תמונות",
            filetypes=[("Images", "*.png *.jpg *.jpeg *.webp *.bmp"), ("All files", "*.*")]
        )
        if not paths:
            return

        for p in paths:
            self.items.append({"path": p, "x": "", "y": ""})
            self.tree.insert("", "end", values=(os.path.basename(p), "", ""))

    def delete_selected(self):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        del self.items[idx]
        self.tree.delete(sel[0])
        self.x_var.set("")
        self.y_var.set("")
        self.preview_label.config(image="")
        self.preview_label.image = None
        self.caption_label.config(text="")

    def on_select(self, _evt=None):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        item = self.items[idx]
        self.x_var.set(item["x"])
        self.y_var.set(item["y"])
        self.show_preview(item["path"], item["x"], item["y"])

    def save_xy(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning("שגיאה", "בחר שורה קודם")
            return

        idx = self.tree.index(sel[0])
        x = self.x_var.get().strip()
        y = self.y_var.get().strip()

        self.items[idx]["x"] = x
        self.items[idx]["y"] = y

        self.tree.item(sel[0], values=(os.path.basename(self.items[idx]["path"]), x, y))
        self.show_preview(self.items[idx]["path"], x, y)

    def show_preview(self, path, x, y):
        try:
            img = Image.open(path)
            img.thumbnail((340, 460))
            tkimg = ImageTk.PhotoImage(img)
            self.preview_label.config(image=tkimg)
            self.preview_label.image = tkimg
        except Exception:
            self.preview_label.config(text="לא ניתן להציג תמונה")

        self.caption_label.config(text=make_caption(x, y) if x and y else "")

    def _validate_items(self):
        if not self.items:
            messagebox.showwarning("שגיאה", "לא הוספת תמונות")
            return False

        missing = []
        for i, it in enumerate(self.items, start=1):
            if not it["x"] or not it["y"]:
                missing.append(i)

        if missing:
            messagebox.showwarning("שגיאה", f"חסרים ערכים X/Y בשורות: {missing}")
            return False
        return True

    def export_pdf(self):
        if not self._validate_items():
            return

        out = filedialog.asksaveasfilename(
            title="שמור PDF",
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")]
        )
        if not out:
            return

        try:
            self._export_pdf_impl(out)
            messagebox.showinfo("הצלחה", f"נשמר PDF:\n{out}")
        except Exception as e:
            messagebox.showerror("שגיאה", f"נכשל ייצוא PDF:\n{e}")

    def _export_pdf_impl(self, out_path):
        c = canvas.Canvas(out_path, pagesize=A4)
        w, h = A4

        margin_x = 15 * mm
        margin_top = 15 * mm
        margin_bottom = 15 * mm

        caption_h = 10 * mm
        gap_between_blocks = 10 * mm

        available_h = h - margin_top - margin_bottom
        block_h = (available_h - gap_between_blocks) / 2.0
        img_h = block_h - caption_h
        img_w = w - 2 * margin_x

        c.setFont("Helvetica", 12)

        for idx, it in enumerate(self.items):
            if idx % 2 == 0:
                if idx != 0:
                    c.showPage()

            block_index = idx % 2
            y_top = h - margin_top - block_index * (block_h + gap_between_blocks)

            img_y = y_top - img_h
            img_x = margin_x

            img = Image.open(it["path"])
            iw, ih = img.size
            scale = min(img_w / iw, img_h / ih)
            dw = iw * scale
            dh = ih * scale

            dx = img_x + (img_w - dw) / 2
            dy = img_y + (img_h - dh) / 2

            c.drawImage(ImageReader(img), dx, dy, width=dw, height=dh,
                        preserveAspectRatio=True, mask='auto')

            caption = make_caption(it["x"], it["y"])
            cap_y = img_y - 6
            c.drawRightString(w - margin_x, cap_y, caption)

        c.save()

    def export_docx(self):
        if not self._validate_items():
            return

        out = filedialog.asksaveasfilename(
            title="שמור DOCX",
            defaultextension=".docx",
            filetypes=[("Word DOCX", "*.docx")]
        )
        if not out:
            return

        try:
            self._export_docx_impl(out)
            messagebox.showinfo("הצלחה", f"נשמר DOCX:\n{out}")
        except Exception as e:
            messagebox.showerror("שגיאה", f"נכשל ייצוא DOCX:\n{e}")

    def _export_docx_impl(self, out_path):
        doc = Document()

        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

        usable_width = section.page_width - section.left_margin - section.right_margin
        img_width_inches = usable_width / Inches(1)

        for i, it in enumerate(self.items):
            if i > 0 and i % 2 == 0:
                doc.add_page_break()

            doc.add_picture(it["path"], width=Inches(img_width_inches))

            p = doc.add_paragraph(make_caption(it["x"], it["y"]))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(12)

            doc.add_paragraph("")

        doc.save(out_path)


if __name__ == "__main__":
    app = App()
    app.mainloop()
